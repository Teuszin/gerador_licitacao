import docx
import locale
import re
import datetime
import pandas as pd
from num2words import num2words

# Importações específicas do docx para formatação
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# =============================================================================
# PARTE 2: FUNÇÕES AUXILIARES (DE AMBOS OS CÓDIGOS)
# =============================================================================

# --- Função do Código 1: Substituir texto no documento ---

def formatar_linha_total(cells, font_size=Pt(5)):
    """Aplica formatação padrão para uma linha de totais na tabela."""
    # Formata a célula de rótulo (que foi mesclada)
    label_cell = cells[0]
    label_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_label = label_cell.paragraphs[0]
    p_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if len(p_label.runs) > 0:
        run_label = p_label.runs[0]
        run_label.bold = True
        run_label.font.size = font_size
        
    # Formata as células de valor
    # Começa do índice 3 porque as células 0, 1, 2 foram mescladas
    for i in range(3, len(cells)):
        value_cell = cells[i]
        value_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_value = value_cell.paragraphs[0]
        p_value.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if len(p_value.runs) > 0:
            run_value = p_value.runs[0]
            run_value.bold = True
            run_value.font.size = font_size


def docx_replace_regex(doc_obj, regex, replace, font_name='Arial', font_size_pt=6.5):
    """
    Encontra e substitui texto usando expressão regular, lidando com 'split runs'
    e aplicando a formatação de fonte especificada.
    """
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            full_text = "".join(r.text for r in inline)
            if regex.search(full_text):
                new_text = regex.sub(replace, full_text)
                for i in range(len(inline)):
                    p.runs[i].text = ""
                new_run = p.add_run(new_text)
                new_run.font.name = font_name
                new_run.font.size = Pt(font_size_pt)

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if regex.search(p.text):
                        inline = p.runs
                        full_text = "".join(r.text for r in inline)
                        if regex.search(full_text):
                            new_text = regex.sub(replace, full_text)
                            for i in range(len(inline)):
                                p.runs[i].text = ""
                            new_run = p.add_run(new_text)
                            new_run.font.name = font_name
                            new_run.font.size = Pt(font_size_pt)


def formatar_porcentagem(valor):
    """Formata um número como uma string de porcentagem com duas casas decimais."""
    return f"{valor:.2f}%".replace('.', ',')

# --- Funções do Código 2: Formatação da planilha ---


def formatar_moeda(valor):
    """Formata um número para o padrão R$ 1.234,56"""
    if pd.isna(valor):
        return ""
    # Usa a configuração de locale já definida no script principal
    return locale.currency(valor, grouping=True)


def set_cell_shading(cell, fill_color):
    """Aplica uma cor de fundo a uma célula da tabela."""
    shading_elm = parse_xml(
        r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), fill_color.replace("#", "")))
    cell._tc.get_or_add_tcPr().append(shading_elm)


# --- FUNÇÃO PRINCIPAL DO CÓDIGO 2 (MODIFICADA) ---
def adicionar_planilha_ao_documento(documento, df_itens, dados_cabecalho):
    """
    Adiciona a planilha orçamentária formatada a um OBJETO DE DOCUMENTO EXISTENTE.
    A função foi modificada para não criar um novo documento nem salvá-lo.
    """
    print("Iniciando a criação da planilha no documento existente...")
    print("\n--- Verificando colunas em 'adicionar_planilha_ao_documento' ---")
    print(df_itens)

    # Define o estilo padrão para esta seção (opcional, mas bom para consistência)
    style = documento.styles['Normal']
    font = style.font
    font.name = 'Arial'

    # --- ALTERAÇÃO: Tabela agora com 10 colunas ---
    tabela = documento.add_table(rows=7, cols=10)
    tabela.style = 'Table Grid'

    # --- FASE 1: Construção do Cabeçalho Estático ---
    celula_titulo = tabela.cell(0, 0).merge(tabela.cell(0, 9))
    celula_titulo.text = "PLANILHA ORÇAMENTÁRIA - ORÇAMENTO SINTÉTICO"
    set_cell_shading(celula_titulo, '#d8ecf6')
    celula_titulo.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    celula_titulo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    run_titulo = celula_titulo.paragraphs[0].runs[0]
    run_titulo.font.bold = True
    run_titulo.font.size = Pt(5)
    tabela.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    tabela.rows[0].height = Cm(0.85)

    # Preenchimento do restante do cabeçalho estático...
    texto_pregao = f"Pregão Eletrônico nº {dados_cabecalho.get('PREGAO', '')} - {dados_cabecalho.get('CLIENTE', '')}"

    tabela.cell(1, 0).merge(tabela.cell(1, 6)).text = texto_pregao
    tabela.cell(1, 0).merge(tabela.cell(1, 6)
                            ).paragraphs[0].runs[0].font.size = Pt(5)

    tabela.cell(1, 7).text = "SINAPI"
    tabela.cell(1, 8).text = "12/2024"
    tabela.cell(1, 9).text = "29,92%"
    tabela.cell(2, 0).merge(tabela.cell(2, 6))
    tabela.cell(2, 7).merge(tabela.cell(2, 8)).text = "Desonerado"
    tabela.cell(2, 9).text = "BDI"

    celula_objeto = tabela.cell(3, 0).merge(tabela.cell(5, 6))
    celula_objeto.text = f"Objeto: {dados_cabecalho.get('OBJETO', '')}"

    celula_objeto.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    celula_objeto.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    celula_objeto.paragraphs[0].runs[0].font.bold = True
    celula_objeto.paragraphs[0].runs[0].font.size = Pt(5)
    tabela.cell(3, 7).merge(tabela.cell(3, 8)).text = "Pernambuco"
    tabela.cell(4, 7).text = "SICRO3"
    tabela.cell(4, 8).text = "08/07/2025"  # Data de hoje como exemplo
    tabela.cell(5, 7).merge(tabela.cell(5, 9))

    for row_idx in range(1, 6):
        for col_idx in range(7, 10):
            cell = tabela.cell(row_idx, col_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if len(cell.text.strip()) > 0:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].runs[0].font.size = Pt(5)

    # --- FASE 2: Títulos da Planilha ---
    titulos = ["Item", "Código", "Banco", "Descrição", "Und", "Quant", "Valor Unit",
               "Valor Unit com BDI", "Valor Total sem BDI", "Valor Total com BDI"]
    cor_fundo_titulo = "#4d93d9"
    cor_letra_titulo = RGBColor(0xFF, 0xFF, 0xFF)

    for i, texto_titulo in enumerate(titulos):
        celula = tabela.cell(6, i)
        paragrafo = celula.paragraphs[0]
        paragrafo.text = ""  # Limpa qualquer texto pré-existente
        run = paragrafo.add_run(texto_titulo)
        run.font.color.rgb = cor_letra_titulo
        run.font.bold = True
        run.font.size = Pt(5)
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        celula.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(celula, cor_fundo_titulo)
    tabela.rows[6].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    tabela.rows[6].height = Cm(1.0)

    # --- FASE 3: Itens dinâmicos da Planilha ---
    for _, row_data in df_itens.iterrows():
        row_cells = tabela.add_row().cells
        row_cells[3].text = row_data['Descrição']
        row_cells[6].text = formatar_moeda(row_data['Valor Unit'])
        row_cells[7].text = formatar_moeda(row_data['Valor Unit com BDI'])
        row_cells[8].text = formatar_moeda(row_data['Valor Total sem BDI'])
        row_cells[9].text = formatar_moeda(row_data['Valor Total com BDI'])
        for i in range(6, 10):
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        row_cells[0].text = str(row_data['Item'])
        row_cells[1].text = str(row_data['Código'])
        row_cells[2].text = str(row_data['Banco'])
        row_cells[4].text = str(row_data['Und'])
        row_cells[5].text = str(row_data['Quant'])
        for i in [0, 1, 2, 4, 5]:
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for cell in row_cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(5)

    # --- FASE 4: Linha de TOTAL ---

    # <<< LÓGICA ADICIONADA AQUI >>>
    # Criamos um novo DataFrame que contém apenas as linhas que SÃO itens
    # (ou seja, onde 'Valor Unit' não é zero).
    # Isso exclui as linhas de "tópico" das nossas somas.
    print("Filtrando apenas os itens válidos para a soma final...")
    df_apenas_itens = df_itens[df_itens['Valor Unit'] > 0]
    # <<< FIM DA LÓGICA ADICIONADA >>>

    # Agora, todas as somas usarão 'df_apenas_itens' em vez de 'df_itens'

    # --- Linha de Total Sem BDI ---
    # <<< MODIFICADO >>>
    valor_total_soma = df_apenas_itens['Valor Total sem BDI'].sum()
    total_row_cells = tabela.add_row().cells

    total_row_cells[0].merge(total_row_cells[7])
    total_label_cell = total_row_cells[8]
    p_total_label = total_label_cell.paragraphs[0]
    run_total_label = p_total_label.add_run('Total Sem BDI')
    # ... (o resto da formatação da célula continua igual) ...
    run_total_label.bold = True
    run_total_label.font.size = Pt(5)
    p_total_label.alignment = WD_ALIGN_PARAGRAPH.LEFT

    total_value_cell = total_row_cells[9]
    p_total_value = total_value_cell.paragraphs[0]
    run_total_value = p_total_value.add_run(formatar_moeda(valor_total_soma))
    # ... (o resto da formatação da célula continua igual) ...
    run_total_value.bold = True
    run_total_value.font.size = Pt(5)
    p_total_value.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    ##################################################

    # --- Linha de Total do BDI ---
    # <<< MODIFICADO >>>
    valor_total_com_bdi = df_apenas_itens['Valor Total com BDI'].sum()
    valor_total_sem_bdi = df_apenas_itens['Valor Total sem BDI'].sum()
    soma_dos_totais = valor_total_com_bdi - valor_total_sem_bdi

    total_row_cells = tabela.add_row().cells
    total_row_cells[0].merge(total_row_cells[7])
    total_label_cell = total_row_cells[8]
    p_total_label = total_label_cell.paragraphs[0]
    run_total_label = p_total_label.add_run('Total do BDI')
    # ... (o resto da formatação da célula continua igual) ...
    run_total_label.bold = True
    run_total_label.font.size = Pt(5)
    p_total_label.alignment = WD_ALIGN_PARAGRAPH.LEFT

    total_value_cell = total_row_cells[9]
    p_total_value = total_value_cell.paragraphs[0]
    run_total_value = p_total_value.add_run(formatar_moeda(soma_dos_totais))
    # ... (o resto da formatação da célula continua igual) ...
    run_total_value.bold = True
    run_total_value.font.size = Pt(5)
    p_total_value.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    ################################

    # --- Linha de Total Geral ---
    # <<< MODIFICADO >>>
    valor_total_soma = df_apenas_itens['Valor Total com BDI'].sum()
    total_row_cells = tabela.add_row().cells

    total_row_cells[0].merge(total_row_cells[7])
    total_label_cell = total_row_cells[8]
    p_total_label = total_label_cell.paragraphs[0]
    run_total_label = p_total_label.add_run('Total Geral')
    # ... (o resto da formatação da célula continua igual) ...
    run_total_label.bold = True
    run_total_label.font.size = Pt(5)
    p_total_label.alignment = WD_ALIGN_PARAGRAPH.LEFT

    total_value_cell = total_row_cells[9]
    p_total_value = total_value_cell.paragraphs[0]
    run_total_value = p_total_value.add_run(formatar_moeda(valor_total_soma))
    # ... (o resto da formatação da célula continua igual) ...
    run_total_value.bold = True
    run_total_value.font.size = Pt(5)
    p_total_value.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    print("Linhas de totais calculadas e adicionadas corretamente.")

    print("Iniciando a criação da planilha no documento existente...")
    print("Planilha adicionada ao documento.")


def adicionar_cronograma_fisico_financeiro(documento, df_cronograma, df_itens, dados_cabecalho, valor_total_projeto):
    """
    Adiciona a planilha de Cronograma Físico-Financeiro formatada a um OBJETO DE DOCUMENTO EXISTENTE.
    Esta função cria uma tabela com 2 linhas por item de dados.

    Args:
        documento: O objeto de documento docx existente.
        df_cronograma: DataFrame do Pandas com os dados do cronograma. 
                       Deve conter colunas para 'Item', 'Descrição', e valores/percentuais para cada etapa.
                       Ex: ['Item', 'Descrição', 'Total_Etapa_Percent', '30_Dias_Percent', ..., 
                            'Total_Etapa_Valor', '30_Dias_Valor', ...]
        df_itens: O DataFrame original dos itens, usado para calcular o valor total geral com BDI.
        dados_cabecalho: Dicionário com informações como PREGAO, CLIENTE e OBJETIVO.
    """
    print("Iniciando a criação do Cronograma Físico-Financeiro...")
    print("\n--- Verificando colunas em 'adicionar_cronograma_fisico_financeiro' ---")
    print(df_itens.columns)  # df_itens é o DataFrame do orçamento aqui também

    # Define o estilo padrão
    style = documento.styles['Normal']
    font = style.font
    font.name = 'Arial'

    # --- Tabela agora com 7 colunas ---
    colunas_tabela = df_cronograma.columns.tolist()
    tabela = documento.add_table(rows=4, cols=len(colunas_tabela))
    tabela.style = 'Table Grid'

    # --- FASE 1: Construção do Cabeçalho Estático ---
    # Linha 1: Título
    celula_titulo = tabela.cell(0, 0).merge(
        tabela.cell(0, len(colunas_tabela)-1))
    celula_titulo.text = "CRONOGRAMA FÍSICO E FINANCEIRO"
    # Mesma cor do cabeçalho da outra planilha
    set_cell_shading(celula_titulo, '#d8ecf6')
    celula_titulo.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    celula_titulo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    run_titulo = celula_titulo.paragraphs[0].runs[0]
    run_titulo.font.bold = True
    run_titulo.font.size = Pt(5)

    # Linha 2: Pregão e Cliente
    celula_pregao = tabela.cell(1, 0).merge(
        tabela.cell(1, len(colunas_tabela)-1))
    texto_pregao = f"Pregão Eletrônico nº {dados_cabecalho.get('PREGAO', '')} - {dados_cabecalho.get('CLIENTE', '')}"
    celula_pregao.text = texto_pregao
    celula_pregao.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    celula_pregao.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    celula_pregao.paragraphs[0].runs[0].font.size = Pt(5)

    # Linha 3: Objeto
    celula_objeto = tabela.cell(2, 0).merge(
        tabela.cell(2, len(colunas_tabela)-1))
    texto_objeto = f"Objeto: {dados_cabecalho.get('OBJETO', '')}"
    celula_objeto.text = texto_objeto
    celula_objeto.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    celula_objeto.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    celula_objeto.paragraphs[0].runs[0].font.size = Pt(5)
    

    for i in range(3):
        tabela.rows[i].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        tabela.rows[i].height = Cm(0.8)

    # --- FASE 2: Títulos da Planilha ---
    cor_fundo_titulo = "#4d93d9"
    cor_letra_titulo = RGBColor(0xFF, 0xFF, 0xFF)

    celulas_titulo = tabela.rows[3].cells
    for i, texto_titulo in enumerate(colunas_tabela):
        celula = celulas_titulo[i]
        paragrafo = celula.paragraphs[0]
        paragrafo.text = ""  # Limpa texto pré-existente
        run = paragrafo.add_run(texto_titulo)
        run.font.color.rgb = cor_letra_titulo
        run.font.bold = True
        run.font.size = Pt(6)
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        celula.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(celula, cor_fundo_titulo)
    tabela.rows[3].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    tabela.rows[3].height = Cm(1.0)

    # =======================================================================
    # FASE 3: Itens dinâmicos da Planilha (Lógica 1-para-1)
    # =======================================================================
    # Esta nova versão é mais simples: ela percorre cada linha do seu
    # DataFrame e cria uma linha correspondente na tabela do Word.

    # 'colunas_tabela' deve ser definida na FASE 1 ou 2, assim:
    # colunas_tabela = df_cronograma.columns.tolist()

    # --- FASE 3: Itens dinâmicos da Planilha ---
    for _, row_data in df_cronograma.iterrows():
        row_cells = tabela.add_row().cells

        # Preenche todas as células primeiro
        for i, col_name in enumerate(colunas_tabela):
            cell_text = str(row_data[col_name])
            row_cells[i].text = cell_text

        # <<< INÍCIO DA CORREÇÃO DE FORMATAÇÃO >>>
        # Agora, percorre as células da linha recém-criada para aplicar a formatação
        for i, cell in enumerate(row_cells):
            # Centraliza verticalmente em todas as células
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            paragrafo = cell.paragraphs[0]
            
            # Alinhamento horizontal: Descrição à esquerda, resto centralizado
            if i == 1: # Coluna "Descrição" (índice 1)
                paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Ajusta o tamanho da fonte em todas as células
            if len(paragrafo.runs) > 0:
                run = paragrafo.runs[0]
                run.font.size = Pt(5)
        # <<< FIM DA CORREÇÃO DE FORMATAÇÃO >>>

    # =======================================================================
    # <<< FASE 4: Linhas de TOTAL (LÓGICA REFEITA) >>>
    # =======================================================================
    # Como a mesma coluna agora tem R$ e %, precisamos ser espertos para somar.

    # 1. Filtramos apenas as linhas que contêm valores financeiros (R$).
    # Usamos a coluna 'Total da Etapa' como referência.
    df_financeiro = df_cronograma[df_cronograma['Total da Etapa'].astype(
        str).str.strip().str.startswith('R$')].copy()

    # 2. Criamos uma função para limpar os valores (remover "R$", espaços, trocar ",")
    def limpar_e_converter(series):
        limpo = series.astype(str).str.replace(
            'R$', '', regex=False).str.strip()
        limpo = limpo.str.replace(
            '.', '', regex=False).str.replace(',', '.', regex=False)
        return pd.to_numeric(limpo, errors='coerce').fillna(0)

    # 3. Calculamos as somas para cada etapa
    soma_etapa1 = limpar_e_converter(df_financeiro['30 Dias']).sum()
    soma_etapa2 = limpar_e_converter(df_financeiro['60 Dias']).sum()
    soma_etapa3 = limpar_e_converter(df_financeiro['90 Dias']).sum()
    soma_etapa4 = limpar_e_converter(df_financeiro['120 Dias']).sum()

    print("\n--- VALORES DE DEBUG (SOMAS) ---")
    print(f"DEBUG: Soma Etapa 1 = {soma_etapa1}")
    print(f"DEBUG: Soma Etapa 2 = {soma_etapa2}")
    print(f"DEBUG: Soma Etapa 3 = {soma_etapa3}")
    print(f"DEBUG: Soma Etapa 4 = {soma_etapa4}")
    print("------------------------------------\n")

    # valor_total_projeto = df_apenas_itens['Valor Total com BDI'].sum()
    valor_total_geral_com_bdi = valor_total_projeto
    # --- Adiciona as linhas de totais na tabela do Word ---

    # Evitar divisão por zero
    if valor_total_geral_com_bdi == 0:
        pct_etapa1, pct_etapa2, pct_etapa3, pct_etapa4 = 0, 0, 0, 0
    else:
        pct_etapa1 = (soma_etapa1 / valor_total_geral_com_bdi) * 100
        pct_etapa2 = (soma_etapa2 / valor_total_geral_com_bdi) * 100
        pct_etapa3 = (soma_etapa3 / valor_total_geral_com_bdi) * 100
        pct_etapa4 = (soma_etapa4 / valor_total_geral_com_bdi) * 100

    # Linha de Porcentagem Mensal
    pct_cells = tabela.add_row().cells
    pct_cells[0].merge(pct_cells[2]).text = 'Porcentagem Mensal'
    pct_cells[3].text = formatar_porcentagem(pct_etapa1)
    pct_cells[4].text = formatar_porcentagem(pct_etapa2)
    pct_cells[5].text = formatar_porcentagem(pct_etapa3)
    pct_cells[6].text = formatar_porcentagem(pct_etapa4)
    formatar_linha_total(pct_cells, font_size=Pt(5))

    # Linha de Custo Mensal
    custo_cells = tabela.add_row().cells
    custo_cells[0].merge(custo_cells[2]).text = 'Custo Mensal'
    custo_cells[3].text = formatar_moeda(soma_etapa1)
    custo_cells[4].text = formatar_moeda(soma_etapa2)
    custo_cells[5].text = formatar_moeda(soma_etapa3)
    custo_cells[6].text = formatar_moeda(soma_etapa4)
    formatar_linha_total(custo_cells, font_size=Pt(5))


    # Linha de Porcentagem Acumulada
    pct_acum_cells = tabela.add_row().cells
    pct_acum_cells[0].merge(pct_acum_cells[2]).text = 'Porcentagem Acumulada'
    acum_pct_1 = pct_etapa1
    acum_pct_2 = acum_pct_1 + pct_etapa2
    acum_pct_3 = acum_pct_2 + pct_etapa3
    acum_pct_4 = acum_pct_3 + pct_etapa4
    pct_acum_cells[3].text = formatar_porcentagem(acum_pct_1)
    pct_acum_cells[4].text = formatar_porcentagem(acum_pct_2)
    pct_acum_cells[5].text = formatar_porcentagem(acum_pct_3)
    pct_acum_cells[6].text = formatar_porcentagem(acum_pct_4)
    formatar_linha_total(pct_acum_cells, font_size=Pt(5))

    # Linha de Custo Acumulado
    # Linha de Custo Acumulado
    custo_acum_cells = tabela.add_row().cells
    custo_acum_cells[0].merge(custo_acum_cells[2]).text = 'Custo Acumulado'
    acum_1 = soma_etapa1
    acum_2 = acum_1 + soma_etapa2
    acum_3 = acum_2 + soma_etapa3
    acum_4 = acum_3 + soma_etapa4

    print("\n--- VALORES DE DEBUG (ACUMULADOS) ---")
    print(f"DEBUG: Acumulado 1 = {acum_1}")
    print(f"DEBUG: Acumulado 2 = {acum_2}")
    print(f"DEBUG: Acumulado 3 = {acum_3}")
    print(f"DEBUG: Acumulado 4 = {acum_4}")
    print("---------------------------------------\n")

    custo_acum_cells[3].text = formatar_moeda(acum_1)
    custo_acum_cells[4].text = formatar_moeda(acum_2)
    custo_acum_cells[5].text = formatar_moeda(acum_3)
    custo_acum_cells[6].text = formatar_moeda(acum_4)
    formatar_linha_total(custo_acum_cells, font_size=Pt(5))

    # (Adicione formatação para as linhas de totais se desejar)
    print("Iniciando a criação do Cronograma (versão 1-para-1)...")
    print("Cronograma Físico-Financeiro (versão 1-para-1) adicionado ao documento.")
